from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import random
#doc = docx.Document("testfile.docx")
mydoc = Document()
diction = {}
font_styles = mydoc.styles
font_charstyle = font_styles.add_style("QuizStyle", WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(14)
font_object.name = "Times New Roman"
section = mydoc.sections[0]
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '2')
defcount = 0
#numex = 0
rlist = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z', 'AA', 'AB', 'AC', 'AD', 'AE', 'AF', 'AG', 'AH', 'AI', 'AJ', 'AK', 'AL', 'AM', 'AN', 'AO', 'AP', 'AQ', 'AR', 'AS', 'AT', 'AU', 'AV', 'AW', 'AX', 'AY', 'AZ', 'BA', 'BB', 'BC', 'BD', 'BE', 'BF', 'BG']
wordlist = []
temprlist = []
adiction = {}
qname = input("Please Enter Quiz Name: ")
ofile = input("Please Enter txt file location: ")
sfile = input("Please Enter name of docx: ")
mydoc.add_heading(qname, 0)
with open(ofile) as o:
    lines = o.readlines()
for line in lines:
    if ':' in line:
        defcount = defcount + 1
temprlist = rlist[0:defcount]
random.shuffle(temprlist)
for line in lines:
    if ':' in line:
        word = line.split(':')
        diction[word[0]] = word[1]
        adiction[word[0]] = temprlist.pop()



        
first_para = mydoc.add_paragraph("")
for d in diction:
    first_para.add_run("____" + d, style="QuizStyle")
    first_para.add_run("\n")
first_para.add_run("\n\n\n\n\n\n")
for let in rlist:
    for ad in adiction:
        if adiction[ad] is let:
            for d in diction:
                if d == ad:
                    first_para.add_run(adiction[ad] + ". " + diction[d] + "\n", style="QuizStyle")

first_para.add_run("\n\n\n\n\n\n\n\n\n\n")
mydoc.add_heading("Answer Sheet", 0)
second_para = mydoc.add_paragraph("")
for ad in adiction:
    for d in diction:
        if d == ad:
            second_para.add_run("__{0}__ {1}\n".format(adiction[ad], d), style="QuizStyle")









mydoc.save(sfile + ".docx")